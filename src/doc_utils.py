import os
import win32com.client
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# Function to add a Table of Contents (TOC) to the document
def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    # Create a empty XML element for the TOC field
    fldChar = OxmlElement('w:fldChar')
    # Set the field type to 'begin' to indicate the start of the TOC field
    fldChar.set(qn('w:fldCharType'), 'begin')
    # Create an XML element for the instruction text of the TOC field
    instrText = OxmlElement('w:instrText')
    # Preserve whitespace in the instruction text to ensure proper formatting
    instrText.set(qn('xml:space'), 'preserve')
    # Comands for the TOC field: \o "1-6" specifies that the TOC should include headings from levels 1 to 6, \h creates hyperlinks in the TOC, \z hides page numbers in web layout view, and \u uses the applied paragraph outline levels to build the TOC  
    instrText.text = 'TOC \\o "1-6" \\h \\z \\u'
    # Create a second XML element 
    fldChar2 = OxmlElement('w:fldChar')
    # Set the field type to 'separate' to indicate the separation between the field instructions and the field result
    fldChar2.set(qn('w:fldCharType'), 'separate')
    # Create a third XML element for the end of the TOC field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Function to add a uniform heading to the document with specified text and level
def add_uniform_heading(document, text, level):
    heading_level = min(level, 6)
    para = document.add_heading('', level=heading_level)
    run = para.add_run(text)

    h1 = document.styles['Heading 1']
    if h1.font.name:
        run.font.name = h1.font.name
    if h1.font.size:
        run.font.size = h1.font.size
    run.font.bold = True if h1.font.bold is None else h1.font.bold

    run.font.italic = False
    try:
        if h1.font.color.rgb:
            run.font.color.rgb = h1.font.color.rgb
    except Exception:
        pass

    return para

# Function to update the Table of Contents (TOC) in the generated document
def update_toc(output_path):
    abs_path = os.path.abspath(output_path)
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(abs_path)
    toc = doc.TablesOfContents(1)
    toc.Update()
    for para in toc.Range.Paragraphs:
        para.Range.Bold = False
        para.Range.Font.Name = 'Arial'
    doc.Save()
    doc.Close()
    word.Quit()